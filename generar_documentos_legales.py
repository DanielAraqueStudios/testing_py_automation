#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de Documentos para Registro Legal de Software en Colombia
Sistema de Cotizaciones Automatizado
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from datetime import datetime
import shutil
import zipfile

def create_legal_registration_folder():
    """Crear la estructura de carpetas para el registro legal"""
    base_path = Path(__file__).parent
    legal_folder = base_path / "registro_legal_colombia"
    
    # Crear carpeta principal
    legal_folder.mkdir(exist_ok=True)
    
    # Crear subcarpetas
    docs_folder = legal_folder / "documentos"
    assets_folder = legal_folder / "recursos"
    
    docs_folder.mkdir(exist_ok=True)
    assets_folder.mkdir(exist_ok=True)
    
    print(f"✅ Carpeta creada: {legal_folder}")
    return legal_folder, docs_folder, assets_folder

def add_table_borders(table):
    """Agregar bordes a la tabla de Word"""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Agregar bordes
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)

def create_program_description_word(docs_folder):
    """Crear la descripción detallada del programa en formato WORD"""
    print("📄 Generando descripción del programa (WORD)...")
    
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título principal
    title = doc.add_heading('DESCRIPCIÓN DETALLADA DEL PROGRAMA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(102, 34, 204)  # Color #6622CC
    
    # Subtítulo
    subtitle = doc.add_heading('Sistema de Cotizaciones Automatizado', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.color.rgb = RGBColor(102, 34, 204)
    
    doc.add_paragraph()
    
    # 1. INFORMACIÓN GENERAL
    doc.add_heading('1. INFORMACIÓN GENERAL DEL SOFTWARE', level=1)
    
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Nombre del Software:', 'Sistema de Cotizaciones Automatizado'),
        ('Versión:', '2.0.0'),
        ('Tipo de Software:', 'Aplicación de Escritorio con Interfaz Gráfica'),
        ('Lenguaje de Programación:', 'Python 3.11+'),
        ('Framework GUI:', 'PyQt6'),
        ('Plataforma:', 'Windows, Linux, macOS'),
        ('Desarrollador:', 'Daniel Araque Studios'),
        ('Fecha de Desarrollo:', datetime.now().strftime('%B %Y'))
    ]
    
    for i, (key, value) in enumerate(info_data):
        info_table.cell(i, 0).text = key
        info_table.cell(i, 1).text = value
        # Hacer negrita la primera columna
        info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    add_table_borders(info_table)
    
    # 2. PROPÓSITO Y FUNCIONALIDAD
    doc.add_heading('2. PROPÓSITO Y FUNCIONALIDAD', level=1)
    
    purpose_text = """
El Sistema de Cotizaciones Automatizado es una aplicación de software especializada diseñada para automatizar y optimizar el proceso completo de generación de propuestas comerciales para empresas del sector tecnológico, específicamente aquellas que ofrecen servicios de:

• Desarrollo web y comercio electrónico
• Marketing digital y redes sociales  
• Automatización con bots de WhatsApp
• Campañas de publicidad pagada
• Capacitación en inteligencia artificial

El software transforma un proceso manual que tradicionalmente requiere varias horas de trabajo en una tarea automatizada que se completa en minutos, manteniendo altos estándares de profesionalismo y consistencia visual.
    """
    doc.add_paragraph(purpose_text)
    
    # 3. DESCRIPCIÓN TÉCNICA DETALLADA
    doc.add_heading('3. DESCRIPCIÓN TÉCNICA DETALLADA', level=1)
    
    # 3.1 Arquitectura
    doc.add_heading('3.1 Arquitectura del Software', level=2)
    architecture_text = """
El software implementa una arquitectura modular basada en el patrón MVC (Modelo-Vista-Controlador):

MODELO (Backend - generate_quote.py):
• Lógica de negocio para procesamiento de datos
• Algoritmos de cálculo de precios y descuentos
• Generación de templates HTML
• Gestión de archivos y estructura de directorios

VISTA (Frontend - quote_gui.py):
• Interfaz gráfica desarrollada en PyQt6
• Formularios de captura de datos del cliente
• Sistema de validación en tiempo real
• Controles de selección de servicios y precios

CONTROLADOR:
• Coordinación entre la interfaz y la lógica de negocio
• Manejo de eventos del usuario
• Procesamiento asíncrono con threading
• Gestión de errores y retroalimentación al usuario
    """
    doc.add_paragraph(architecture_text)
    
    # 3.2 Componentes principales
    doc.add_heading('3.2 Componentes Principales', level=2)
    
    components_text = """
MÓDULO DE INTERFAZ GRÁFICA:
• Clase ModernQuoteApp: Ventana principal de la aplicación
• Clase PriceInputWidget: Widgets especializados para entrada de precios
• Sistema de pestañas para organización de datos
• Validación automática de campos de entrada

MÓDULO DE PROCESAMIENTO:
• Función get_user_input(): Captura de datos del usuario
• Función select_services(): Selección de servicios a cotizar
• Función parse_price_input(): Validación y formateo de precios
• Función generate_client_folder_structure(): Creación de estructura de archivos

MÓDULO DE GENERACIÓN:
• Sistema de templates HTML con sustitución de variables
• Generación automática de carpetas por cliente
• Copia automática de recursos (CSS, JS, imágenes)
• Creación de estructura lista para web hosting

MÓDULO DE UTILIDADES:
• Formateo de precios con separadores de miles colombianos
• Sanitización de nombres de archivos y carpetas
• Validación de datos de entrada
• Manejo de errores y logging
    """
    doc.add_paragraph(components_text)
    
    # 4. ALGORITMOS Y PROCEDIMIENTOS
    doc.add_heading('4. ALGORITMOS Y PROCEDIMIENTOS PRINCIPALES', level=1)
    
    # 4.1 Algoritmo de cálculo de precios
    doc.add_heading('4.1 Algoritmo de Cálculo de Precios con Descuento', level=2)
    
    price_algorithm = """
ENTRADA: 
- precio_final (precio que paga el cliente)
- porcentaje_descuento (0-99%)

PROCESO:
1. Validar que precio_final > 0
2. Validar que 0 ≤ porcentaje_descuento < 100
3. Calcular multiplicador = 1 - (porcentaje_descuento / 100)
4. Calcular precio_original = precio_final / multiplicador
5. Formatear números con separadores de miles colombianos

SALIDA:
- precio_original (precio antes del descuento)
- precio_final (precio con descuento aplicado)

EJEMPLO:
Entrada: precio_final = 4.500.000, descuento = 10%
Proceso: multiplicador = 1 - (10/100) = 0.9
         precio_original = 4.500.000 / 0.9 = 5.000.000
Salida: Original: $5.000.000, Final: $4.500.000
    """
    doc.add_paragraph(price_algorithm)
    
    # 4.2 Algoritmo de generación de estructura
    doc.add_heading('4.2 Algoritmo de Generación de Estructura de Archivos', level=2)
    
    structure_algorithm = """
ENTRADA:
- nombre_cliente (string)
- datos_cotizacion (diccionario con información)
- template_html (archivo base)

PROCESO:
1. Sanitizar nombre_cliente para nombre de carpeta válido
2. Crear ruta: cotizaciones/[nombre_cliente]/to_upload/
3. Copiar carpeta assets/ completa al destino
4. Procesar template HTML con datos de la cotizacion:
   a. Reemplazar variables de cliente y empresa
   b. Aplicar secciones de servicios seleccionados
   c. Insertar precios formateados
   d. Configurar términos y condiciones
5. Guardar como index.html en la carpeta to_upload/
6. Verificar integridad de archivos generados

SALIDA:
- Estructura de carpetas lista para deployment
- Archivo index.html funcional
- Recursos completos (CSS, JS, imágenes)
- Mensaje de confirmación con rutas generadas
    """
    doc.add_paragraph(structure_algorithm)
    
    # 5. ESPECIFICACIONES TÉCNICAS
    doc.add_heading('5. ESPECIFICACIONES TÉCNICAS', level=1)
    
    # Tabla de especificaciones
    specs_table = doc.add_table(rows=10, cols=2)
    specs_table.style = 'Table Grid'
    
    specs_data = [
        ('Lenguaje Base:', 'Python 3.11+'),
        ('Framework GUI:', 'PyQt6'),
        ('Librerías Adicionales:', 'python-docx, reportlab, shutil, pathlib'),
        ('Formato de Salida:', 'HTML5 + CSS3 + JavaScript'),
        ('Compatibilidad:', 'Windows 10+, macOS 10.14+, Linux Ubuntu 18+'),
        ('Memoria RAM Mínima:', '4 GB'),
        ('Espacio en Disco:', '100 MB + espacio para cotizaciones'),
        ('Resolución de Pantalla:', '1024x768 mínimo (recomendado 1920x1080)'),
        ('Conexión de Red:', 'No requerida para funcionamiento básico'),
        ('Formato de Datos:', 'UTF-8, números con formato colombiano')
    ]
    
    for i, (key, value) in enumerate(specs_data):
        specs_table.cell(i, 0).text = key
        specs_table.cell(i, 1).text = value
        specs_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    add_table_borders(specs_table)
    
    # 6. PROCEDIMIENTO DE INSTALACIÓN
    doc.add_heading('6. PROCEDIMIENTO DE INSTALACIÓN Y CONFIGURACIÓN', level=1)
    
    installation_text = """
REQUISITOS PREVIOS:
1. Sistema operativo compatible (Windows 10+, macOS 10.14+, Linux Ubuntu 18+)
2. Espacio en disco: mínimo 100 MB
3. Memoria RAM: mínimo 4 GB

PASOS DE INSTALACIÓN (VERSIÓN EJECUTABLE):
1. Descargar el ejecutable del software (empaquetado con PyInstaller)
2. Extraer archivos en directorio deseado
3. Ejecutar el archivo Sistema_Cotizaciones.exe (Windows) o ejecutable correspondiente
4. El software está listo para usar sin instalaciones adicionales

NOTA: El software viene empaquetado con PyInstaller e incluye todas las dependencias necesarias.
No requiere instalación previa de Python ni librerías adicionales.

CONFIGURACIÓN INICIAL:
1. Configurar datos de empresa por defecto
2. Personalizar templates HTML según branding
3. Ajustar precios base de servicios
4. Configurar términos y condiciones estándar

VERIFICACIÓN DE FUNCIONAMIENTO:
1. Generar cotización de prueba
2. Verificar creación de estructura de archivos
3. Comprobar visualización de HTML generado
4. Validar integridad de recursos (CSS, JS, imágenes)
    """
    doc.add_paragraph(installation_text)
    
    # 7. CASOS DE USO
    doc.add_heading('7. CASOS DE USO PRINCIPALES', level=1)
    
    use_cases = """
CASO DE USO 1: Generación de Cotización Completa
Actor: Usuario (Comercial/Vendedor)
Flujo: Ingresar datos → Seleccionar servicios → Configurar precios → Generar
Resultado: Carpeta completa con cotización lista para envío

CASO DE USO 2: Cálculo Automático de Precios
Actor: Usuario
Flujo: Ingresar precio final → Seleccionar % descuento → Sistema calcula precio original
Resultado: Precios originales y finales calculados automáticamente

CASO DE USO 3: Personalización de Servicios
Actor: Usuario
Flujo: Seleccionar servicios específicos → Solo aparecen secciones relevantes
Resultado: Cotización personalizada solo con servicios seleccionados

CASO DE USO 4: Generación Masiva
Actor: Usuario avanzado
Flujo: Procesar múltiples clientes → Generar estructura organizada
Resultado: Múltiples carpetas organizadas en /cotizaciones/
    """
    doc.add_paragraph(use_cases)
    
    # 8. MARCO LEGAL Y NORMATIVO
    doc.add_heading('8. MARCO LEGAL Y NORMATIVO', level=1)
    
    legal_framework = """
CUMPLIMIENTO DE REQUISITOS LEGALES EN COLOMBIA

Este software cumple con todos los requisitos establecidos para el registro legal de software en Colombia:

1. PRODUCCIÓN DE SOFTWARE
   • Código fuente completo en formato comprimido
   • Algoritmos e instrucciones según el lenguaje de programación (Python)
   • Manual técnico y programa ejecutable incluidos
   • Documentación técnica detallada

2. RESOLUCIÓN 00285 DE COLCIENCIAS
   Expedida el 19 de marzo de 2004
   
   El software satisface los criterios establecidos en esta resolución:
   • Innovación tecnológica en automatización de procesos comerciales
   • Desarrollo de algoritmos propietarios para cálculo de precios
   • Implementación de interfaz gráfica moderna
   • Generación automática de contenido web
   • Valor agregado en optimización de procesos empresariales

3. REGISTRO DE DERECHOS DE AUTOR
   • Entidad: Ministerio del Interior - Dirección Nacional de Derechos de Autor
   • Tipo: Certificado de Registro de Soporte Lógico
   • Propósito: Protección legal del software desarrollado
   • Estado: Preparado para solicitud de certificado

DOCUMENTACIÓN INCLUIDA PARA REGISTRO:
✓ Descripción detallada del programa (este documento)
✓ Material auxiliar para usuarios
✓ Código fuente comprimido con algoritmos
✓ Manuales técnicos y de instalación
✓ Especificaciones técnicas completas
    """
    doc.add_paragraph(legal_framework)
    
    # 9. CLASIFICACIÓN SEGÚN RESOLUCIÓN 00285
    doc.add_heading('9. CLASIFICACIÓN DEL SOFTWARE SEGÚN RESOLUCIÓN 00285', level=1)
    
    classification = """
TIPO DE PRODUCCIÓN: PRODUCCIÓN TECNOLÓGICA

Según la Resolución 00285 de Colciencias, este software se clasifica como PRODUCCIÓN TECNOLÓGICA porque genera innovación al:

• Desarrollar una nueva aplicación especializada en automatización de cotizaciones
• Implementar mejoras significativas sobre procesos manuales existentes
• Crear algoritmos propietarios de cálculo automático de precios con descuentos
• Extender herramientas existentes (PyQt6) con funcionalidad especializada
• Mejorar significativamente la metodología de generación de propuestas comerciales
• Innovar en la integración de cálculo financiero con generación automática de web

El software NO se clasifica como producción científica porque no genera conocimiento teórico nuevo (nuevos modelos algorítmicos, nuevas metodologías de desarrollo, nuevos sistemas operativos, etc.), sino que aplica tecnologías existentes de manera innovadora para resolver un problema específico del sector comercial.
    """
    doc.add_paragraph(classification)
    
    # 10. CRITERIOS TÉCNICOS DE PONDERACIÓN (RESOLUCIÓN 00285)
    doc.add_heading('10. CRITERIOS TÉCNICOS DE PONDERACIÓN', level=1)
    
    doc.add_paragraph('Evaluación del software según los criterios técnicos establecidos en la Resolución 00285 de Colciencias:')
    doc.add_paragraph()
    
    # Crear tabla de criterios técnicos
    criteria_table = doc.add_table(rows=10, cols=2)
    criteria_table.style = 'Table Grid'
    
    criteria_data = [
        ('CRITERIO', 'CUMPLIMIENTO EN ESTE SOFTWARE'),
        ('1. ROBUSTEZ', 'El software es sólido y funciona correctamente incluso con datos inesperados. Implementa validación exhaustiva de entradas, manejo de errores con try-except, y recuperación graciosa ante fallos. No se bloquea ante situaciones difíciles.'),
        ('2. EXTENDIBILIDAD', 'Arquitectura modular basada en MVC facilita agregar nuevas funcionalidades. Nuevos servicios, métodos de cálculo o formatos de salida pueden agregarse sin modificar el código existente. Código bien estructurado y documentado.'),
        ('3. DESEMPEÑO', 'Cumple tareas en tiempo real sin demoras perceptibles. Genera cotizaciones completas en segundos. Uso eficiente de memoria RAM (<50MB). No desperdicia espacio en disco. Threading para operaciones largas.'),
        ('4. USABILIDAD', 'Interfaz gráfica intuitiva diseñada para usuarios no técnicos. Organización en pestañas lógicas. Validación en tiempo real. Mensajes claros. Documentación completa incluida. Curva de aprendizaje mínima.'),
        ('5. INTEGRIDAD', 'Datos consistentes en todo el proceso. Validación de entradas previene corrupción. Archivos generados mantienen estructura íntegra. No hay pérdida de información. Sistema de verificación post-generación.'),
        ('6. PORTABILIDAD', 'Compatible con Windows, macOS y Linux sin modificaciones. Python multiplataforma garantiza portabilidad. Rutas de archivos con pathlib funcionan en todos los sistemas. Sin dependencias específicas de SO.'),
        ('7. COMPATIBILIDAD', 'Versión 2.0.0 mantiene compatibilidad con estructuras de datos de versión 1.x. Templates HTML son retrocompatibles. Archivos generados funcionan en cualquier navegador moderno.'),
        ('8. MANTENIMIENTO', 'Código limpio y bien documentado facilita mantenimiento. Arquitectura modular permite actualizaciones aisladas. Comentarios explican lógica compleja. Estructura de proyecto organizada lógicamente.'),
        ('9. DOCUMENTACIÓN', 'Completamente documentado con descripciones técnicas, manuales de usuario, código fuente comentado, README detallado, y documentación para registro legal. Cumple estándares de documentación profesional.')
    ]
    
    for i, (criterion, compliance) in enumerate(criteria_data):
        criteria_table.cell(i, 0).text = criterion
        criteria_table.cell(i, 1).text = compliance
        # Hacer negrita la primera fila (encabezados)
        if i == 0:
            criteria_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
            criteria_table.cell(i, 1).paragraphs[0].runs[0].font.bold = True
        else:
            criteria_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    add_table_borders(criteria_table)
    
    doc.add_paragraph()
    conclusion = """
CONCLUSIÓN: El software cumple satisfactoriamente con los 9 criterios técnicos de ponderación establecidos en la Resolución 00285 de Colciencias, demostrando ser un producto tecnológico robusto, bien diseñado, documentado y apto para registro legal.
    """
    conclusion_para = doc.add_paragraph(conclusion)
    conclusion_para.runs[0].font.bold = True
    
    # Footer con información legal
    doc.add_paragraph("\n" * 3)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"Documento generado el {datetime.now().strftime('%d de %B de %Y')}")
    footer_run.font.size = Pt(10)
    footer_run.italic = True
    
    footer_para2 = doc.add_paragraph()
    footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run2 = footer_para2.add_run("Daniel Araque Studios - Registro Legal de Software Colombia")
    footer_run2.font.size = Pt(10)
    footer_run2.italic = True
    
    # Guardar documento
    word_path = docs_folder / "Descripcion_Detallada_Programa.docx"
    doc.save(word_path)
    print(f"✅ Descripción del programa guardada: {word_path}")
    return word_path

def create_auxiliary_material_pdf(docs_folder):
    """Crear el material auxiliar en formato PDF"""
    print("📄 Generando material auxiliar (PDF)...")
    
    pdf_path = docs_folder / "Material_Auxiliar_Usuario.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
    
    # Crear estilos
    styles = getSampleStyleSheet()
    
    # Estilo personalizado para títulos
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.Color(102/255, 34/255, 204/255)  # #6622CC
    )
    
    # Estilo para subtítulos
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=20,
        textColor=colors.Color(102/255, 34/255, 204/255)
    )
    
    # Estilo para texto normal
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_JUSTIFY
    )
    
    # Contenido del PDF
    content = []
    
    # Título principal
    content.append(Paragraph("MATERIAL AUXILIAR", title_style))
    content.append(Paragraph("Sistema de Cotizaciones Automatizado", subtitle_style))
    content.append(Spacer(1, 20))
    
    # Introducción
    content.append(Paragraph("INTRODUCCIÓN", subtitle_style))
    intro_text = """
Este material auxiliar proporciona información complementaria para comprender y utilizar 
eficientemente el Sistema de Cotizaciones Automatizado. El documento incluye guías de uso, 
explicaciones de funcionalidades, ejemplos prácticos y consejos para optimizar el flujo 
de trabajo en la generación de propuestas comerciales.
    """
    content.append(Paragraph(intro_text, normal_style))
    content.append(Spacer(1, 20))
    
    # Guía de inicio rápido
    content.append(Paragraph("GUÍA DE INICIO RÁPIDO", subtitle_style))
    
    quick_start = """
<b>1. PREPARACIÓN INICIAL</b><br/>
• Descargue el ejecutable del software (empaquetado con PyInstaller)<br/>
• Extraiga el contenido en una carpeta de su preferencia<br/>
• No requiere instalación de Python ni dependencias adicionales<br/><br/>

<b>2. PRIMER USO</b><br/>
• Ejecute el archivo Sistema_Cotizaciones.exe (Windows) o el ejecutable correspondiente<br/>
• Complete la pestaña "Información" con datos del cliente<br/>
• Seleccione servicios en la pestaña "Servicios"<br/>
• Configure precios en la pestaña "Precios"<br/>
• Haga clic en "Generar Cotización"<br/><br/>

<b>3. VERIFICACIÓN</b><br/>
• Revise la carpeta /cotizaciones/[nombre_cliente]/<br/>
• Abra el archivo index.html en un navegador<br/>
• Verifique que todos los recursos carguen correctamente<br/>
    """
    content.append(Paragraph(quick_start, normal_style))
    content.append(Spacer(1, 20))
    
    # Explicación de la interfaz
    content.append(Paragraph("EXPLICACIÓN DE LA INTERFAZ", subtitle_style))
    
    interface_text = """
<b>PESTAÑA INFORMACIÓN:</b><br/>
- Datos del Cliente: Nombres, cargo, empresa<br/>
- Datos del Remitente: Su información como vendedor<br/>
- Términos: Días de validez y condiciones de pago<br/><br/>

<b>PESTAÑA SERVICIOS:</b><br/>
- Desarrollo Web: Sitios ecommerce profesionales<br/>
- Redes Sociales: Estrategias de marketing digital<br/>
- Bot WhatsApp: Automatización de atención al cliente<br/>
- Campañas Facebook: Publicidad pagada optimizada<br/>
- Capacitación IA: Formación en herramientas de inteligencia artificial<br/><br/>

<b>PESTAÑA PRECIOS:</b><br/>
- Precio Final: Lo que pagará el cliente<br/>
- % Descuento: Porcentaje de descuento aplicado<br/>
- Precio Original: Calculado automáticamente por el sistema<br/>
    """
    content.append(Paragraph(interface_text, normal_style))
    content.append(Spacer(1, 20))
    
    # Casos de uso prácticos
    content.append(Paragraph("CASOS DE USO PRÁCTICOS", subtitle_style))
    
    use_cases_text = """
<b>CASO PRÁCTICO 1: Cotización Web Básica</b><br/>
Cliente: "Restaurante La Parrilla"<br/>
Servicios: Solo desarrollo web<br/>
Precio final: $4.500.000<br/>
Descuento: 10%<br/>
Resultado: Precio original $5.000.000 calculado automáticamente<br/><br/>

<b>CASO PRÁCTICO 2: Paquete Completo</b><br/>
Cliente: "Boutique Fashion Style"<br/>
Servicios: Web + Redes Sociales + Bot WhatsApp<br/>
Estructura: Precios individuales por cada servicio<br/>
Resultado: Cotización completa con múltiples secciones<br/><br/>

<b>CASO PRÁCTICO 3: Solo Marketing Digital</b><br/>
Cliente: "Consultorio Médico Salud+"<br/>
Servicios: Redes Sociales + Campañas Facebook<br/>
Enfoque: Solo servicios de marketing<br/>
Resultado: Cotización especializada en marketing<br/>
    """
    content.append(Paragraph(use_cases_text, normal_style))
    content.append(PageBreak())
    
    # Solución de problemas
    content.append(Paragraph("SOLUCIÓN DE PROBLEMAS COMUNES", subtitle_style))
    
    troubleshooting_text = """
<b>PROBLEMA: "El ejecutable no inicia"</b><br/>
Solución: Verificar que extrajo todos los archivos del paquete completo.<br/>
Asegurar permisos de ejecución. En Windows, permitir la aplicación si Windows Defender la bloquea.<br/><br/>

<b>PROBLEMA: "Los campos de precio no permiten números grandes"</b><br/>
Solución: El sistema permite hasta 20 dígitos. Escriba sin puntos ni comas,
el sistema formateará automáticamente.<br/><br/>

<b>PROBLEMA: "No se genera la carpeta del cliente"</b><br/>
Soluciones:<br/>
- Verificar permisos de escritura en el directorio<br/>
- Asegurar que el nombre del cliente no tenga caracteres especiales<br/>
- Revisar que todos los campos obligatorios estén completos<br/><br/>

<b>PROBLEMA: "El HTML no se ve correctamente"</b><br/>
Soluciones:<br/>
- Verificar que la carpeta /assets/ se copió correctamente<br/>
- Abrir el index.html desde la carpeta /to_upload/<br/>
- Verificar conexión a internet para fuentes externas<br/>
    """
    content.append(Paragraph(troubleshooting_text, normal_style))
    content.append(Spacer(1, 20))
    
    # Consejos de optimización
    content.append(Paragraph("CONSEJOS DE OPTIMIZACIÓN", subtitle_style))
    
    optimization_text = """
<b>PARA MAYOR EFICIENCIA:</b><br/>
• Configure datos de empresa por defecto para no repetir información<br/>
• Use nombres de cliente descriptivos para mejor organización<br/>
• Mantenga una estructura consistente en términos de pago<br/>
• Revise periódicamente la carpeta /cotizaciones/ para limpieza<br/><br/>

<b>PARA MEJOR PRESENTACIÓN:</b><br/>
• Personalice el archivo template.html con su branding<br/>
• Ajuste colores corporativos en los archivos CSS<br/>
• Incluya su logo en la carpeta /assets/images/<br/>
• Adapte textos descriptivos a su modelo de negocio<br/><br/>

<b>PARA MAYOR PRODUCTIVIDAD:</b><br/>
• Use el modo claro/oscuro según su preferencia<br/>
• Aproveche el cálculo automático de precios<br/>
• Genere múltiples cotizaciones para comparación<br/>
• Mantenga backup de cotizaciones importantes<br/>
    """
    content.append(Paragraph(optimization_text, normal_style))
    content.append(Spacer(1, 20))
    
    # Glosario de términos
    content.append(Paragraph("GLOSARIO DE TÉRMINOS", subtitle_style))
    
    # Crear tabla para el glosario
    glossary_data = [
        ['Término', 'Definición'],
        ['Template', 'Plantilla HTML base para generar cotizaciones'],
        ['Assets', 'Recursos como CSS, JavaScript e imágenes'],
        ['Backend', 'Lógica de procesamiento del software'],
        ['Frontend', 'Interfaz gráfica visible al usuario'],
        ['Threading', 'Procesamiento en segundo plano'],
        ['Sanitización', 'Limpieza de caracteres especiales en nombres'],
        ['Deployment', 'Preparación de archivos para publicación web'],
        ['Responsive', 'Diseño adaptable a diferentes pantallas'],
        ['GUI', 'Interfaz Gráfica de Usuario']
    ]
    
    glossary_table = Table(glossary_data, colWidths=[2*inch, 4*inch])
    glossary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(102/255, 34/255, 204/255)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    
    content.append(glossary_table)
    content.append(Spacer(1, 30))
    
    # Información de contacto y soporte
    content.append(Paragraph("INFORMACIÓN DE SOPORTE", subtitle_style))
    
    contact_text = """
<b>DESARROLLADOR:</b> Daniel Araque Studios<br/>
<b>VERSIÓN:</b> 2.0.0<br/>
<b>FECHA DE CREACIÓN:</b> """ + datetime.now().strftime('%B %Y') + """<br/>
<b>SISTEMA OPERATIVO:</b> Windows, macOS, Linux<br/><br/>

<b>SOPORTE TÉCNICO:</b><br/>
Para consultas técnicas, actualizaciones o personalizaciones del software,
contacte al desarrollador a través de los canales oficiales.<br/><br/>

<b>ACTUALIZACIONES:</b><br/>
Se recomienda mantener el software actualizado para acceder a nuevas
funcionalidades y mejoras de seguridad.<br/>
    """
    content.append(Paragraph(contact_text, normal_style))
    
    # Footer
    content.append(Spacer(1, 40))
    footer_text = f"Material Auxiliar generado automáticamente el {datetime.now().strftime('%d de %B de %Y')}"
    footer_para = Paragraph(footer_text, ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        alignment=TA_CENTER,
        textColor=colors.grey
    ))
    content.append(footer_para)
    
    # Generar PDF
    doc.build(content)
    print(f"✅ Material auxiliar guardado: {pdf_path}")
    return pdf_path

def create_readme_file(legal_folder):
    """Crear archivo README con información de los documentos"""
    readme_path = legal_folder / "README.md"
    
    readme_content = f"""# Registro Legal del Software en Colombia

## Sistema de Cotizaciones Automatizado

### Documentos Generados para Registro Legal

Este directorio contiene la documentación requerida para el registro legal del software en Colombia según la normativa vigente.

#### 📁 Estructura de Archivos

```
registro_legal_colombia/
├── documentos/
│   ├── Descripcion_Detallada_Programa.docx  (Descripción técnica en WORD)
│   └── Material_Auxiliar_Usuario.pdf        (Material auxiliar en PDF)
├── recursos/
│   └── (archivos de apoyo)
└── README.md (este archivo)
```

#### 📋 Documentos Incluidos

**1. Descripción Detallada del Programa (WORD)**
- Información general del software
- Arquitectura y componentes técnicos
- Algoritmos y procedimientos principales
- Especificaciones técnicas completas
- Casos de uso y procedimientos de instalación

**2. Material Auxiliar (PDF)**
- Guía de inicio rápido
- Explicación de la interfaz de usuario
- Casos de uso prácticos
- Solución de problemas comunes
- Consejos de optimización
- Glosario de términos técnicos

#### 🎯 Propósito del Software

El Sistema de Cotizaciones Automatizado es una aplicación especializada para:
- Automatización de propuestas comerciales
- Cálculo automático de precios y descuentos
- Generación de estructuras web listas para deployment
- Optimización del proceso de ventas

#### 📊 Especificaciones Técnicas

- **Lenguaje:** Python 3.11+
- **Framework GUI:** PyQt6
- **Plataformas:** Windows, macOS, Linux
- **Tipo:** Software de escritorio con interfaz gráfica
- **Versión:** 2.0.0

#### 📅 Información de Registro

- **Fecha de generación:** {datetime.now().strftime('%d de %B de %Y')}
- **Desarrollador:** Daniel Araque Studios
- **Propósito:** Registro legal en Colombia
- **Estado:** Listo para presentación

#### 📞 Contacto

Para consultas sobre el registro legal o el software:
- Desarrollador: Daniel Araque Studios
- Fecha: {datetime.now().strftime('%Y')}

---
*Documentos generados automáticamente para cumplir con los requisitos de registro de software en Colombia*
"""
    
    with open(readme_path, 'w', encoding='utf-8') as f:
        f.write(readme_content)
    
    print(f"✅ README creado: {readme_path}")
    return readme_path

def create_source_code_package(docs_folder):
    """Crear paquete comprimido con código fuente y algoritmos"""
    print("📦 Creando paquete de código fuente comprimido...")
    
    base_path = Path(__file__).parent
    zip_path = docs_folder / "Codigo_Fuente_Comprimido.zip"
    
    # Archivos a incluir en el paquete
    files_to_include = [
        'generate_quote.py',
        'quote_gui.py',
        'run_gui.py',
        'template.html',
        'generar_documentos_legales.py',
        'README.md'
    ]
    
    # Crear archivo ZIP
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        # Agregar archivos Python (código fuente)
        for filename in files_to_include:
            file_path = base_path / filename
            if file_path.exists():
                zipf.write(file_path, filename)
                print(f"   ✓ Agregado: {filename}")
        
        # Agregar carpeta assets si existe
        assets_path = base_path / 'assets'
        if assets_path.exists():
            for root, dirs, files in os.walk(assets_path):
                for file in files:
                    file_path = Path(root) / file
                    arcname = file_path.relative_to(base_path)
                    zipf.write(file_path, arcname)
        
        # Crear archivo LEEME.txt con instrucciones
        readme_content = """CÓDIGO FUENTE - Sistema de Cotizaciones Automatizado
===============================================================

CONTENIDO DEL PAQUETE:
- generate_quote.py: Lógica principal y algoritmos de generación
- quote_gui.py: Interfaz gráfica de usuario (PyQt6)
- run_gui.py: Ejecutable principal
- template.html: Plantilla HTML base
- generar_documentos_legales.py: Generador de documentación legal
- assets/: Recursos web (CSS, JavaScript, imágenes)

LENGUAJE DE PROGRAMACIÓN: Python 3.11+

INSTRUCCIONES DE USO:
El software se distribuye como ejecutable empaquetado con PyInstaller.
No requiere instalación de Python ni dependencias adicionales.

1. Extraer el contenido del paquete
2. Ejecutar el archivo Sistema_Cotizaciones.exe (Windows) o ejecutable correspondiente
3. El software está listo para usar inmediatamente

NOTA: PyInstaller empaqueta Python y todas las librerías necesarias en un ejecutable standalone.

ALGORITMOS PRINCIPALES:
- Cálculo automático de precios con descuento
- Generación de estructura de archivos web
- Procesamiento de templates HTML con sustitución de variables
- Validación y sanitización de datos de entrada

FRAMEWORK GUI: PyQt6
COMPATIBILIDAD: Windows, macOS, Linux

REGISTRO LEGAL:
Este código fuente está incluido en el registro legal del software
ante el Ministerio del Interior - Dirección Nacional de Derechos de Autor
según Resolución 00285 de Colciencias (19 de marzo de 2004).

DESARROLLADOR: Daniel Araque Studios
VERSIÓN: 2.0.0
FECHA: """ + datetime.now().strftime('%B %Y') + """

"""
        zipf.writestr('LEEME.txt', readme_content)
        
        # Crear archivo con especificaciones técnicas
        specs_content = """ESPECIFICACIONES TÉCNICAS
==========================

REQUISITOS DEL SISTEMA:
- Sistema Operativo: Windows 10+, macOS 10.14+, Linux Ubuntu 18+
- RAM: 4 GB mínimo
- Espacio en disco: 150 MB + espacio para cotizaciones
- Resolución: 1024x768 mínimo
- NO requiere instalación previa de Python

DEPENDENCIAS (INCLUIDAS EN EL EJECUTABLE):
El software se distribuye empaquetado con PyInstaller e incluye:
- PyQt6: Framework de interfaz gráfica
- python-docx: Generación de documentos WORD
- reportlab: Generación de documentos PDF
- Python 3.11+ runtime
- Todas las librerías necesarias

NOTA: El usuario final NO necesita instalar nada adicional.

ARQUITECTURA:
Patrón MVC (Modelo-Vista-Controlador)
- Modelo: generate_quote.py
- Vista: quote_gui.py  
- Controlador: Coordinación entre módulos

FORMATO DE SALIDA:
HTML5 + CSS3 + JavaScript (Bootstrap, jQuery)

CODIFICACIÓN: UTF-8
FORMATO DE NÚMEROS: Separadores de miles colombianos

"""
        zipf.writestr('ESPECIFICACIONES_TECNICAS.txt', specs_content)
    
    print(f"✅ Paquete de código fuente creado: {zip_path}")
    return zip_path

def main():
    """Función principal para generar todos los documentos legales"""
    print("🚀 Iniciando generación de documentos para registro legal...")
    print("=" * 60)
    
    try:
        # Crear estructura de carpetas
        legal_folder, docs_folder, assets_folder = create_legal_registration_folder()
        
        # Generar documentos
        word_doc = create_program_description_word(docs_folder)
        pdf_doc = create_auxiliary_material_pdf(docs_folder)
        source_code_zip = create_source_code_package(docs_folder)
        readme_doc = create_readme_file(legal_folder)
        
        # Resumen final
        print("\n" + "=" * 60)
        print("✅ DOCUMENTOS GENERADOS EXITOSAMENTE")
        print("=" * 60)
        print(f"📁 Carpeta principal: {legal_folder}")
        print(f"📄 Descripción WORD: {word_doc.name}")
        print(f"📄 Material auxiliar PDF: {pdf_doc.name}")
        print(f"📦 Código fuente ZIP: {source_code_zip.name}")
        print(f"📄 README: {readme_doc.name}")
        print("\n⚖️  CUMPLIMIENTO LEGAL:")
        print("   ✓ Resolución 00285 de Colciencias (19 de marzo de 2004)")
        print("   ✓ Código fuente comprimido con algoritmos")
        print("   ✓ Manual técnico y documentación auxiliar")
        print("   ✓ Listo para Certificado del Ministerio del Interior")
        print("\n🎯 Los documentos están listos para el registro legal del software en Colombia")
        print("📋 Revise los archivos generados antes de presentarlos")
        
    except Exception as e:
        print(f"❌ Error al generar documentos: {str(e)}")
        import traceback
        traceback.print_exc()
        return False
    
    return True

if __name__ == "__main__":
    main()